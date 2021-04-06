import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import os

link = "https://www.warandpeace.ru/ru/news/_/_/page=1/"

content = requests.get(link).text
soup = BeautifulSoup(content, 'lxml') # получение кода страницы

pages = soup.findAll("a", class_ = "menu_1") # получение числа страниц новостей
lastPage = "1"
prevTag = "0"
for tag in pages:
    if tag.text == ">":
        lastPage = prevTag # получение номера последней страницы
        break
    prevTag = tag.text

document = Document()
listName = list()
listTime = list()
listHref = list()
listArticles = list()
listTopicsAndSources = list()

with open ("file.html", "w", encoding="utf-8") as file:
    for i in range(1, int(lastPage)+1):
        link = "https://www.warandpeace.ru/ru/news/_/_/page=" + str(i) + "/"
        content = requests.get(link).text
        soup = BeautifulSoup(content, 'lxml')
        block = soup.findAll("a", class_ = "a_header_article") # получение ссылок и названий статей
        date = soup.findAll("td", class_ = "topic_info_top") # получение времени публикации статьи
        for i in range(15):
            href = soup.findAll("a",  class_ = "a_header_article")[i].get('href') # получение ссылки на статью
            listHref.append(href) # запись ссылки на статью в список
        for i in range(1, 31):
            topicAndSource = soup.findAll("a",  class_ = "a_default")[i] # получение темы и источника статьи
            listTopicsAndSources.append(topicAndSource.text) # запись темы и источника статьи в список
        for tag in block:
            listName.append(tag.text) # запись названий статей в список
        for time in date:
            listTime.append(time.text) # запись времени публикации статьи в список

    for i in range(len(listHref)):
        linkArt = listHref[i]
        contentArt = requests.get(linkArt).text
        soupArt = BeautifulSoup(contentArt, 'lxml')  # получение кода страницы статьи
        blockArt = soupArt.findAll(class_="topic_text")  # получение текста статей
        for article in blockArt:
            listArticles.append(article.text) # запись текста статей в список

style = document.styles['Normal'] # создаю стиль основного текста
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

title = document.add_heading("Новости с сайта warandpeace.ru", 0) # создаю заголовок 0
titleStyle0 = title.style # создаю стиль заголовка 0
titleStyle0.font.color.rgb = RGBColor(0, 0, 0) # делаю цвет текста заголовка 0 черным
titleStyle0.font.size = Pt(20) # задаю размер шрифта заголовка 0
title.alignment = WD_ALIGN_PARAGRAPH.CENTER # размещаю заголовок 0 по центру

for i in range(len(listName)):
    nameArticle = document.add_heading(listName[i], 1) # создаю заголовок 1
    titleStyle1 = nameArticle.style # создаю стиль заголовка 1
    titleStyle1.font.color.rgb = RGBColor(0, 0, 0) # делаю текст заголовока 1 черным
    titleStyle1.font.size = Pt(16) # задаю размер шрифта заголовка 1
    nameArticle.alignment = WD_ALIGN_PARAGRAPH.CENTER # размещаю заголовок 1 по центру
    t = document.add_paragraph("\n" + listTime[i]) # добавляю время публикации и размещаю справа
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    top = document.add_paragraph() # добавляю абзац под рубрику публикации и размещаю справа
    top.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    runTop = top.add_run("Рубрика: " + listTopicsAndSources[2*i]) # добавляю рубрику
    runTop.italic = True # делаю текст рубрики курсивом
    b = document.add_paragraph(listArticles[i]) # добавляю текст публикации и размещаю во всю ширину
    b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fb = b.paragraph_format
    fb.first_line_indent = Inches(0.5) # добаляю красную строку
    s = document.add_paragraph() # добавляю абзац под источник публикации и размещаю справа
    s.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    runSource = s.add_run("Источник: " + listTopicsAndSources[2 * i + 1]) # добавляю источник
    runSource.italic = True # делаю источник курсивом
    l = document.add_paragraph("Ссылка: " + listHref[i]) # добавляю ссылку на публикацию и размещаю справа
    l.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = l.add_run()
    run.add_break(WD_BREAK.PAGE) # добавляю разрыв страницы
document.save("News of warandpeace.docx")
os.startfile("News of warandpeace.docx")
